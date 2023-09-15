# Legend
# 1 - Company1 name
# 2 - Company1 address
# 3 - Company2 name
# 4 - Company2 business
# 5 - dd
# 6 - Month name
# 7 - yyyy
# 8 - Year of comapnies act
# 9 - Goods

# hardcoding for temporary purposes (or maybe for permanant)
cache = ["National Payment Corporation of India", "New Parsi house, Lalu Seth lane, Bhandup (W), Mumbai", "Aero & Co.", "Computer Accelerators", "18th", "January", "2023", "1986", "CPUs and GPUs"]


from docx import Document

# creating new output docx
doc = Document()

doc = Document("test.docx")

def replace_string(old, new):
    
    for p in doc.paragraphs:
        if old in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if old in inline[i].text:
                    text = inline[i].text.replace(old, new)
                    inline[i].text = text
            # printp.text
    return 1

for i in range(len(cache)):
    replace_string('#' + str(i+1), cache[i])
      
doc.save("Output2.docx")