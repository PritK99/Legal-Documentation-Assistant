# Legend
# 1 - Company1 name
# 2 - Company1 address
# 3 - Company2 name
# 4 - Company2 business
# 5 - dd
# 6 - Month name
# 7 - yyyy
# 8 - Year of comapnies act
# 9 - Goods

# hardcoding for temporary purposes (or maybe for permanant)
cache = ["National Payment Corporation of India", "New Parsi house, Lalu Seth lane, Bhandup (W), Mumbai", "Aero & Co.", "Computer Accelerators", "18th", "January", "2023", "1986", "CPUs and GPUs"]

from docx import Document

# creating new output docx
doc = Document()

# reading the template
document = Document('test.docx')

# traverse through the document
for paragraph in document.paragraphs:
    temp = paragraph.text
    newstr =""
    maxctr = len(temp)
    i=0
    while(i < maxctr):
        if (temp[i] == '#'):
            index = ""
            i+=1
            while (temp[i] >= '1' and temp[i] <= '9'):
                index+=temp[i]
                i+=1
            i-=1
            newstr+=cache[int(index)-1]
        else:
            newstr+=temp[i]
        i+=1
    print(newstr)
    doc.add_paragraph(newstr)
doc.save("Output.docx")