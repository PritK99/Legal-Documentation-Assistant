# Legend
# 1 - Company1 name
# 2 - Company1 address
# 3 - Company2 name
# 4 - Company2 business
# 5 - dd
# 6 - Month name
# 7 - yyyy
# 8 - Year of comapnies act
# 9 - Goods

# hardcoding for temporary purposes (or maybe for permanant)
cache = ["National Payment Corporation of India", "New Parsi house, Lalu Seth lane, Bhandup (W), Mumbai", "Aero & Co.", "Computer Accelerators", "18th", "January", "2023", "1986", "CPUs and GPUs"]

from docx import Document
import docx2txt

# creating new output docx
doc = Document()

text = docx2txt.process('test.docx')

def insert_hardcoded(x):
    for i in range(0, len(cache)):
        x = x.replace('#' + str(i+1), cache[i])
    print(x)
    return x

# reading the template
# document = Document('test.docx')

# traverse through the document
res = insert_hardcoded(text)
doc.add_paragraph(res)
doc.save("Output2.docx")